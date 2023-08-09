import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
from html2docx import html2docx

import os
import email
import html2text
from docx import Document

def extract_html_from_mht(mht_path):
    with open(mht_path, 'rb') as f:
        msg = email.message_from_binary_file(f)
        html_content = None
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == 'text/html':
                html_content = part.get_payload(decode=True)
                break
        return html_content

def convert_mht_to_word(mht_file):
    # Extract the HTML content from the MHT file
    html_content = extract_html_from_mht(mht_file)

    if not html_content:
        print(f"No HTML content found in {mht_file}")
        return

    # Convert HTML to plain text using html2text (optional)
    plain_text_content = html2text.html2text(html_content.decode('utf-8'))

    # Create a new Word document
    doc = Document()
    doc.add_heading('Converted Document', level=1)

    # Add the plain text content to the Word document
    doc.add_paragraph(plain_text_content)

    # Save the Word document
    output_path = mht_file.replace('.mht', '.docx')
    doc.save(output_path)

    print(f"Converted {mht_file} to Word document: {output_path}")

def handle_dropped_item(event):
    item = event.data.strip('{}')  # Remove curly braces from the path

    if os.path.isdir(item):
        print("Dropped item is a folder. Processing...")
        for root, _, files in os.walk(item):
            for file in files:
                if file.lower().endswith(".mht"):
                    mht_file = os.path.join(root, file)
                    convert_mht_to_word(mht_file)
    elif os.path.isfile(item) and item.lower().endswith(".mht"):
        print("Dropped item is a single MHT file. Processing...")
        convert_mht_to_word(item)
    else:
        # Split and process individual items
        individual_items = item.split('} {')
        for individual_item in individual_items:
            individual_item = individual_item.strip('{}')
            if os.path.isdir(individual_item):
                print(f"Processing folder: {individual_item}")
                for root, _, files in os.walk(individual_item):
                    for file in files:
                        if file.lower().endswith(".mht"):
                            mht_file = os.path.join(root, file)
                            convert_mht_to_word(mht_file)
            elif os.path.isfile(individual_item) and individual_item.lower().endswith(".mht"):
                print(f"Processing MHT file: {individual_item}")
                convert_mht_to_word(individual_item)
            else:
                print(f"Skipping unsupported item: {individual_item}")

root = TkinterDnD.Tk()
root.title("MHT to Word Converter")

frame = tk.Frame(root)
frame.pack(expand=True, fill='both')  # Expand and fill available space

label = tk.Label(frame, text="Drag and drop .MHT file or folder here:")
label.pack(expand=True, fill='both', padx=150, pady=150)  # Expand and fill available space

frame.drop_target_register(DND_FILES)
frame.dnd_bind('<<Drop>>', handle_dropped_item)

root.mainloop()